#!/usr/bin/env python3
"""
album_report — Scan a Photo Albums folder and print a summary of each album:
file counts by type, total size, and rating distribution if a cull_results.json
is present.

Usage:
    album_report "/path/to/Photo Albums"
    album_report "/path/to/Photo Albums" --sort size
    album_report "/path/to/Photo Albums" --save
    album_report "/path/to/Photo Albums" --docx
"""

import argparse
import io
import json
import sys
from collections import Counter
from datetime import date
from pathlib import Path

# ── Config ────────────────────────────────────────────────────────────────────

RAW_EXTENSIONS = {".dng", ".cr2", ".cr3", ".nef", ".arw", ".orf", ".rw2", ".raf"}
JPG_EXTENSIONS = {".jpg", ".jpeg"}
VID_EXTENSIONS = {".mp4", ".mov", ".avi", ".mkv", ".mts", ".m2ts", ".webm"}

# ── Helpers ───────────────────────────────────────────────────────────────────

def human_size(num_bytes: int) -> str:
    for unit in ("B", "KB", "MB", "GB", "TB"):
        if num_bytes < 1024:
            return f"{num_bytes:.1f} {unit}"
        num_bytes /= 1024
    return f"{num_bytes:.1f} PB"


def folder_size(path: Path) -> int:
    return sum(f.stat().st_size for f in path.rglob("*") if f.is_file())


def count_files(path: Path) -> dict:
    counts = {"raw": 0, "jpg": 0, "video": 0, "other": 0}
    if not path.exists():
        return counts
    for f in path.rglob("*"):
        if not f.is_file():
            continue
        ext = f.suffix.lower()
        if ext in RAW_EXTENSIONS:
            counts["raw"] += 1
        elif ext in JPG_EXTENSIONS:
            counts["jpg"] += 1
        elif ext in VID_EXTENSIONS:
            counts["video"] += 1
        else:
            counts["other"] += 1
    return counts


def load_cull_results(album_dir: Path) -> list | None:
    json_path = album_dir / "cull_results.json"
    if not json_path.exists():
        return None
    try:
        with open(json_path) as f:
            return json.load(f)
    except Exception:
        return None


def rating_summary(data: list) -> tuple:
    photos = [r for r in data if r.get("type") in ("raw", "jpg") and "rating" in r]
    if not photos:
        return None, None
    counts  = Counter(r["rating"] for r in photos)
    keepers = sum(counts.get(s, 0) for s in (4, 5))
    total   = len(photos)
    pct     = keepers / total * 100 if total else 0
    stars   = "  ".join(f"*{s}:{counts.get(s, 0)}" for s in range(5, 0, -1))
    return stars, f"{keepers}/{total} ({pct:.0f}%)"


def bar(value: int, maximum: int, width: int = 15) -> str:
    if maximum == 0:
        return "░" * width
    filled = int(width * value / maximum)
    return "█" * filled + "░" * (width - filled)

# ── Scan ──────────────────────────────────────────────────────────────────────

def scan_albums(photo_albums: Path, sort_by: str) -> list:
    if not photo_albums.exists():
        print(f"ERROR: Path not found: {photo_albums}")
        sys.exit(1)
    if not photo_albums.is_dir():
        print(f"ERROR: Not a directory: {photo_albums}")
        sys.exit(1)

    albums = sorted(d for d in photo_albums.iterdir() if d.is_dir())
    if not albums:
        print("No albums found.")
        sys.exit(0)

    results = []
    print(f"Scanning {len(albums)} album(s)...\n")

    for album in albums:
        size   = folder_size(album)
        counts = count_files(album)
        data   = load_cull_results(album)
        rat, keep = rating_summary(data) if data else (None, None)
        results.append({
            "name":            album.name,
            "path":            album,
            "size":            size,
            "counts":          counts,
            "data":            data,
            "rating_summary":  rat,
            "keepers_summary": keep,
        })

    if sort_by == "size":
        results.sort(key=lambda r: r["size"], reverse=True)
    else:
        results.sort(key=lambda r: r["name"].lower())

    return results

# ── Terminal report ───────────────────────────────────────────────────────────

def build_report(results: list, photo_albums: Path) -> str:
    total_size = sum(r["size"] for r in results)
    max_size   = max(r["size"] for r in results) if results else 1
    total_raw  = sum(r["counts"]["raw"]   for r in results)
    total_jpg  = sum(r["counts"]["jpg"]   for r in results)
    total_vid  = sum(r["counts"]["video"] for r in results)

    lines = []
    lines.append("═" * 70)
    lines.append("  PHOTO ALBUMS REPORT")
    lines.append(f"  {photo_albums}")
    lines.append("═" * 70)
    lines.append("")
    lines.append("OVERVIEW")
    lines.append("─" * 70)
    lines.append(f"  Albums      : {len(results)}")
    lines.append(f"  Total size  : {human_size(total_size)}")
    lines.append(f"  RAW files   : {total_raw}")
    lines.append(f"  JPEG files  : {total_jpg}")
    lines.append(f"  Video files : {total_vid}")
    lines.append("")
    lines.append("ALBUMS")
    lines.append("─" * 70)

    for r in results:
        c     = r["counts"]
        total = c["raw"] + c["jpg"] + c["video"]
        lines.append(f"\n  {r['name']}")
        lines.append(f"  {bar(r['size'], max_size)}  {human_size(r['size'])}")
        lines.append(f"  RAW: {c['raw']}  JPEG: {c['jpg']}  Video: {c['video']}  Total: {total}")
        if r["rating_summary"]:
            lines.append(f"  {r['rating_summary']}  keepers: {r['keepers_summary']}")
        else:
            lines.append("  (no cull_results.json)")

    lines.append("")
    lines.append("═" * 70)
    return "\n".join(lines)

# ── Pie chart ─────────────────────────────────────────────────────────────────

def build_pie_chart(results: list) -> tuple:
    """Return (png_bytes, grays) — pie only, no legend (legend goes in docx table)."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    total = sum(r["size"] for r in results)
    sizes = [r["size"] for r in results]
    n     = len(results)

    # Greyscale from dark to light
    grays = [str(round(0.08 + 0.72 * i / max(n - 1, 1), 2)) for i in range(n)]

    fig, ax = plt.subplots(figsize=(5, 5), facecolor="white")

    ax.pie(
        sizes,
        colors=grays,
        startangle=140,
        wedgeprops={"linewidth": 0.8, "edgecolor": "white"},
    )
    ax.set_title(
        f"Storage Distribution  ·  {human_size(total)} total  ·  {n} albums",
        fontsize=10, fontweight="bold", pad=14, color="#111111",
    )

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.read(), grays

# ── DOCX ──────────────────────────────────────────────────────────────────────

def build_docx(results: list, photo_albums: Path, out_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BLACK = RGBColor(0x00, 0x00, 0x00)
    DGRAY = RGBColor(0x33, 0x33, 0x33)
    MGRAY = RGBColor(0x77, 0x77, 0x77)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Remove default paragraph spacing from Normal style
    normal = doc.styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after  = Pt(0)

    # ── Helpers ───────────────────────────────────────────────────────────────

    def add_run(para, text, size=9, bold=False, color=BLACK, font="Courier New"):
        run = para.add_run(text)
        run.font.name      = font
        run.font.size      = Pt(size)
        run.font.bold      = bold
        run.font.color.rgb = color
        return run

    def new_para(space_before=0, space_after=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        return p

    def add_rule(thickness=6, color="999999"):
        p   = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    str(thickness))
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), color)
        pBdr.append(bot)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(3)
        return p

    def shade_cell(cell, fill_hex):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill_hex)
        tcPr.append(shd)

    today = date.today().strftime("%B %d, %Y")

    # ── Wordmark ──────────────────────────────────────────────────────────────
    p = new_para(space_after=1)
    add_run(p, "WebJelly Studios", size=20, bold=True,  color=BLACK, font="Arial")
    add_run(p, "  ·  Photo Library Report", size=12, color=DGRAY, font="Arial")

    p2 = new_para(space_after=5)
    add_run(p2, today, size=8, color=MGRAY, font="Arial")

    add_rule(thickness=14, color="000000")

    p3 = new_para(space_before=4, space_after=8)
    add_run(p3, str(photo_albums), size=7, color=MGRAY)

    # ── Overview ──────────────────────────────────────────────────────────────
    total_size  = sum(r["size"] for r in results)
    total_raw   = sum(r["counts"]["raw"]   for r in results)
    total_jpg   = sum(r["counts"]["jpg"]   for r in results)
    total_vid   = sum(r["counts"]["video"] for r in results)
    total_other = sum(r["counts"]["other"] for r in results)
    total_files = total_raw + total_jpg + total_vid + total_other

    oh = new_para(space_after=3)
    add_run(oh, "OVERVIEW", size=9, bold=True, color=BLACK, font="Arial")
    add_rule(thickness=4, color="999999")

    ov_rows = [
        ("Albums",       str(len(results))),
        ("Total Size",   human_size(total_size)),
        ("RAW Files",    str(total_raw)),
        ("JPEG Files",   str(total_jpg)),
        ("Video Files",  str(total_vid)),
        ("Total Files",  str(total_files)),
    ]
    tbl = doc.add_table(rows=len(ov_rows), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(ov_rows):
        cells = tbl.rows[i].cells
        cells[0].width = Inches(1.5)
        cells[1].width = Inches(2.0)
        if i % 2 == 0:
            shade_cell(cells[0], "F0F0F0")
            shade_cell(cells[1], "F0F0F0")
        lp = cells[0].paragraphs[0]
        vp = cells[1].paragraphs[0]
        lp.paragraph_format.space_before = Pt(1)
        lp.paragraph_format.space_after  = Pt(1)
        vp.paragraph_format.space_before = Pt(1)
        vp.paragraph_format.space_after  = Pt(1)
        add_run(lp, label, size=8, bold=True,  color=DGRAY, font="Arial")
        add_run(vp, value, size=8, bold=False, color=BLACK, font="Courier New")

    new_para(space_after=8)

    # ── Albums ────────────────────────────────────────────────────────────────
    ah = new_para(space_after=3)
    add_run(ah, "ALBUMS", size=9, bold=True, color=BLACK, font="Arial")
    add_rule(thickness=4, color="999999")

    max_size  = max(r["size"] for r in results) if results else 1
    BAR_WIDTH = 28

    for r in results:
        c       = r["counts"]
        total   = c["raw"] + c["jpg"] + c["video"]
        filled  = round(BAR_WIDTH * r["size"] / max_size)
        empty   = BAR_WIDTH - filled
        bar_str = "█" * filled + "░" * empty

        # Name
        np_ = new_para(space_before=6, space_after=1)
        add_run(np_, r["name"], size=9, bold=True, color=BLACK, font="Arial")

        # Bar + size
        bp = new_para(space_after=1)
        add_run(bp, bar_str, size=8, color=DGRAY)
        add_run(bp, f"  {human_size(r['size'])}", size=8, color=MGRAY)

        # Counts
        fp = new_para(space_after=1)
        add_run(fp, f"RAW: {c['raw']}  JPEG: {c['jpg']}  Video: {c['video']}  Other: {c['other']}  Total: {total}",
                size=8, color=DGRAY)

        # Ratings
        rp = new_para(space_after=2)
        if r["rating_summary"]:
            add_run(rp, f"{r['rating_summary']}  keepers: {r['keepers_summary']}", size=8, color=DGRAY)
        else:
            add_run(rp, "(no cull_results.json)", size=8, color=MGRAY)

    add_rule(thickness=4, color="999999")
    new_para(space_after=8)

    # ── Storage distribution ──────────────────────────────────────────────────
    ch = new_para(space_after=3)
    add_run(ch, "STORAGE DISTRIBUTION", size=9, bold=True, color=BLACK, font="Arial")
    add_rule(thickness=4, color="999999")

    png_bytes, grays = build_pie_chart(results)
    img_stream = io.BytesIO(png_bytes)
    doc.add_picture(img_stream, width=Inches(3.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_para(space_after=5)

    # Legend table: swatch | album name | size | %
    total_size_local = sum(r["size"] for r in results)
    leg_tbl = doc.add_table(rows=len(results) + 1, cols=4)
    leg_tbl.style = "Table Grid"
    COL_W = [Inches(0.22), Inches(3.2), Inches(1.1), Inches(0.8)]

    # Header row
    hcells = leg_tbl.rows[0].cells
    for ci, (w, lbl) in enumerate(zip(COL_W, ["", "Album", "Size", "%"])):
        hcells[ci].width = w
        shade_cell(hcells[ci], "E0E0E0")
        hp = hcells[ci].paragraphs[0]
        hp.paragraph_format.space_before = Pt(1)
        hp.paragraph_format.space_after  = Pt(1)
        add_run(hp, lbl, size=7, bold=True, color=BLACK, font="Arial")

    for row_i, (r, gray) in enumerate(zip(results, grays), 1):
        cells = leg_tbl.rows[row_i].cells
        for ci, w in enumerate(COL_W):
            cells[ci].width = w
        if row_i % 2 == 0:
            for ci in range(4):
                shade_cell(cells[ci], "F6F6F6")

        # Swatch cell — filled with gray shade
        gray_hex = format(round(float(gray) * 255), "02X") * 3
        shade_cell(cells[0], gray_hex)
        cells[0].paragraphs[0].add_run(" ")

        pct = r["size"] / total_size_local * 100
        for ci, (txt, fnt) in enumerate([
            (r["name"],             "Arial"),
            (human_size(r["size"]), "Courier New"),
            (f"{pct:.1f}%",         "Courier New"),
        ], 1):
            cp = cells[ci].paragraphs[0]
            cp.paragraph_format.space_before = Pt(1)
            cp.paragraph_format.space_after  = Pt(1)
            add_run(cp, txt, size=7, color=DGRAY, font=fnt)

    new_para(space_after=6)

    # File type mini-bars
    fth = new_para(space_after=2)
    add_run(fth, "File Type Breakdown", size=8, bold=True, color=BLACK, font="Arial")
    MINI = 20
    for label, count in [("RAW", total_raw), ("JPEG", total_jpg), ("Video", total_vid), ("Other", total_other)]:
        p_ = new_para(space_after=1)
        pct    = count / total_files * 100 if total_files else 0
        filled = round(MINI * count / total_files) if total_files else 0
        mini   = "█" * filled + "░" * (MINI - filled)
        add_run(p_, f"  {label:<6}  {mini}  {count:>5} files  ({pct:.1f}%)", size=8, color=DGRAY)

    # ── Footer ────────────────────────────────────────────────────────────────
    new_para(space_after=8)
    add_rule(thickness=10, color="000000")
    fp2 = new_para(space_before=2)
    add_run(fp2, "WebJelly Studios  ·  Confidential", size=7, color=MGRAY, font="Arial")
    add_run(fp2, f"  ·  Generated {today}", size=7, color=MGRAY, font="Arial")

    doc.save(out_path)
    print(f"Report saved → {out_path}")

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Scan a Photo Albums folder and report sizes and file counts"
    )
    parser.add_argument("path", type=str, nargs="?",
                        help="Path to the Photo Albums folder (drag and drop works)")
    parser.add_argument("--sort", default="name", choices=["name", "size"],
                        help="Sort albums by name or size (default: name)")
    parser.add_argument("--save", action="store_true",
                        help="Save plain-text report as album_report.txt on the Desktop")
    parser.add_argument("--docx", action="store_true",
                        help="Generate a WebJelly Studios .docx report on the Desktop")
    args = parser.parse_args()

    raw_path = args.path
    if not raw_path:
        raw_path = input("Photo Albums folder: ").strip().strip("'\"")

    photo_albums = Path(raw_path).expanduser().resolve()
    results      = scan_albums(photo_albums, args.sort)
    report       = build_report(results, photo_albums)
    print(report)

    if args.save:
        out_path = Path.home() / "Desktop" / "album_report.txt"
        out_path.write_text(report)
        print(f"\nReport saved → {out_path}")

    if args.docx:
        out_path = Path.home() / "Desktop" / "album_report.docx"
        build_docx(results, photo_albums, out_path)


if __name__ == "__main__":
    main()
